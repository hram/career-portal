import json
import os
import re
import subprocess
import tempfile
from pathlib import Path

import pdfplumber
from docx import Document
from dotenv import load_dotenv


load_dotenv()

CLAUDE_CLI_PATH = os.getenv("CLAUDE_CLI_PATH", "claude")
CLAUDE_PARSE_TIMEOUT_SECONDS = int(os.getenv("CLAUDE_PARSE_TIMEOUT_SECONDS", "300"))
CLAUDE_STEP_TIMEOUT_SECONDS = int(os.getenv("CLAUDE_STEP_TIMEOUT_SECONDS", "90"))
RESUME_PARSE_PROVIDER = os.getenv("RESUME_PARSE_PROVIDER", "claude-cli")
CODEX_CLI_PATH = os.getenv("CODEX_CLI_PATH", "codex")
CODEX_MODEL = os.getenv("CODEX_MODEL", "")
CODEX_SANDBOX = os.getenv("CODEX_SANDBOX", "read-only")
CODEX_PARSE_TIMEOUT_SECONDS = int(os.getenv("CODEX_PARSE_TIMEOUT_SECONDS", "600"))
CODEX_STEP_TIMEOUT_SECONDS = int(os.getenv("CODEX_STEP_TIMEOUT_SECONDS", "180"))


class ResumeParseError(Exception):
    pass


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        with pdfplumber.open(path) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return _clean_text("\n\n".join(pages))

    if suffix == ".docx":
        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return _clean_text("\n".join(parts))

    raise ResumeParseError("Поддерживаются только PDF и DOCX")


def parse_resume_with_ai(raw_text: str, provider: str | None = None) -> dict:
    if not raw_text.strip():
        raise ResumeParseError("Не удалось извлечь текст из файла")

    selected_provider = (provider or RESUME_PARSE_PROVIDER).strip().lower()
    return parse_resume_pipeline(raw_text, provider=selected_provider)


def parse_resume_pipeline(raw_text: str, provider: str | None = None, on_progress=None) -> dict:
    if not raw_text.strip():
        raise ResumeParseError("Не удалось извлечь текст из файла")

    selected_provider = (provider or RESUME_PARSE_PROVIDER).strip().lower()
    if selected_provider not in {"claude-cli", "codex-cli"}:
        raise ResumeParseError(f"Неизвестный ИИ-агент: {selected_provider}")

    with tempfile.TemporaryDirectory(prefix="career-resume-parse-") as tmpdir:
        resume_text_path = Path(tmpdir) / "resume.txt"
        resume_text_path.write_text(raw_text, encoding="utf-8")
        outline_prompt = (
            _build_outline_file_prompt(resume_text_path)
            if selected_provider == "claude-cli"
            else _build_outline_prompt(raw_text)
        )
        outline_allowed_dir = resume_text_path.parent if selected_provider == "claude-cli" else None

        outline = _run_provider_json(
            selected_provider,
            outline_prompt,
            allowed_dir=outline_allowed_dir,
            timeout=_step_timeout_for_provider(selected_provider),
        )
        result = _normalize_parsed_resume(outline)
        outline_jobs = [item for item in outline.get("jobs") or [] if isinstance(item, dict)]
        result["jobs"] = [_normalize_job(item) for item in outline_jobs]
        total_jobs = len(outline_jobs)

        if on_progress is not None:
            on_progress(
                {
                    "stage": "outline_done",
                    "parsed": result,
                    "total_jobs": total_jobs,
                    "completed_jobs": 0,
                    "current_job": "",
                }
            )

        lines = raw_text.splitlines()
        parsed_jobs: list[dict] = []

        for index, job in enumerate(outline_jobs):
            job_text = _slice_lines(lines, job.get("start_line"), job.get("end_line"))
            if not job_text.strip():
                job_text = "\n".join(value for value in (job.get("company"), job.get("position"), job.get("raw_notes")) if value)

            current_job = job.get("company") or job.get("position") or f"Работа {index + 1}"
            if on_progress is not None:
                on_progress(
                    {
                        "stage": "job_processing",
                        "parsed": {**result, "jobs": [*parsed_jobs, *result["jobs"][index:]]},
                        "total_jobs": total_jobs,
                        "completed_jobs": index,
                        "current_job": current_job,
                    }
                )

            try:
                parsed_job = _run_provider_json(
                    selected_provider,
                    _build_job_prompt(job, job_text),
                    allowed_dir=None,
                    timeout=_step_timeout_for_provider(selected_provider),
                )
            except ResumeParseError:
                parsed_job = job

            normalized_job = _normalize_job({**job, **parsed_job})
            parsed_jobs.append(normalized_job)
            result["jobs"][index] = normalized_job

            if on_progress is not None:
                on_progress(
                    {
                        "stage": "job_done",
                        "parsed": result,
                        "total_jobs": total_jobs,
                        "completed_jobs": index + 1,
                        "current_job": current_job,
                    }
                )

        result["jobs"] = parsed_jobs
        if on_progress is not None:
            on_progress(
                {
                    "stage": "done",
                    "parsed": result,
                    "total_jobs": total_jobs,
                    "completed_jobs": total_jobs,
                    "current_job": "",
                }
            )
        return result


def _step_timeout_for_provider(provider: str) -> int:
    if provider == "codex-cli":
        return CODEX_STEP_TIMEOUT_SECONDS
    return CLAUDE_STEP_TIMEOUT_SECONDS


def _run_provider_json(provider: str, prompt: str, *, allowed_dir: Path | None, timeout: int) -> dict:
    if provider == "claude-cli":
        return _run_claude_json(prompt, allowed_dir=allowed_dir, timeout=timeout)
    if provider == "codex-cli":
        raw_response = _run_codex_text(prompt, timeout=timeout)
        try:
            return _parse_json_response(raw_response)
        except json.JSONDecodeError as exc:
            raise ResumeParseError("Codex CLI вернул невалидный JSON") from exc
    raise ResumeParseError(f"Неизвестный ИИ-агент: {provider}")


def _clean_text(value: str) -> str:
    value = value.replace("\x00", "")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _parse_json_response(raw: str) -> dict:
    cleaned = raw.strip()
    cleaned = re.sub(r"^```json\s*", "", cleaned)
    cleaned = re.sub(r"^```\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start != -1 and end != -1 and end > start:
            return json.loads(cleaned[start : end + 1])
        raise


def _extract_assistant_text(stdout: str) -> list[str]:
    full_text: list[str] = []
    for line in stdout.splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            event = json.loads(line)
        except json.JSONDecodeError:
            continue
        if event.get("type") == "assistant":
            for block in event.get("message", {}).get("content", []):
                if block.get("type") == "text":
                    full_text.append(block.get("text", ""))
    return full_text


def _claude_cli_error_message(completed: subprocess.CompletedProcess[str], raw_response: str = "") -> str | None:
    combined_output = "\n".join(part for part in (raw_response, completed.stderr, completed.stdout) if part)
    if "Not logged in" in combined_output or "Please run /login" in combined_output:
        return "Claude CLI не авторизован: выполните claude /login или настройте ANTHROPIC_API_KEY в окружении"
    if completed.returncode != 0:
        error = (completed.stderr.strip() or raw_response.strip() or "без stderr")[:500]
        return f"Claude CLI завершился с кодом {completed.returncode}: {error}"
    return None


def _run_claude_json(prompt: str, *, allowed_dir: Path | None, timeout: int) -> dict:
    command = [
        CLAUDE_CLI_PATH,
        "-p",
        prompt,
        "--output-format",
        "stream-json",
        "--verbose",
        "--model",
        "sonnet",
        "--permission-mode",
        "bypassPermissions",
    ]
    if allowed_dir is not None:
        command.extend(["--add-dir", str(allowed_dir)])

    try:
        completed = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=timeout,
            check=False,
        )
    except FileNotFoundError as exc:
        raise ResumeParseError(f"Claude CLI не найден: {CLAUDE_CLI_PATH}") from exc
    except subprocess.TimeoutExpired as exc:
        raise ResumeParseError(f"Claude CLI не ответил за {timeout} секунд") from exc

    full_text = _extract_assistant_text(completed.stdout)
    raw_response = "".join(full_text).strip()
    error_message = _claude_cli_error_message(completed, raw_response)
    if error_message:
        raise ResumeParseError(error_message)
    if not raw_response:
        raise ResumeParseError(completed.stderr.strip() or "Claude CLI не вернул ответ")

    try:
        return _parse_json_response(raw_response)
    except json.JSONDecodeError as exc:
        raise ResumeParseError("Claude CLI вернул невалидный JSON") from exc


def _run_claude_text(prompt: str, *, allowed_dir: Path | None, timeout: int) -> str:
    command = [
        CLAUDE_CLI_PATH,
        "-p",
        prompt,
        "--output-format",
        "stream-json",
        "--verbose",
        "--model",
        "sonnet",
        "--permission-mode",
        "bypassPermissions",
    ]
    if allowed_dir is not None:
        command.extend(["--add-dir", str(allowed_dir)])

    try:
        completed = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=timeout,
            check=False,
        )
    except FileNotFoundError as exc:
        raise ResumeParseError(f"Claude CLI не найден: {CLAUDE_CLI_PATH}") from exc
    except subprocess.TimeoutExpired as exc:
        raise ResumeParseError(f"Claude CLI не ответил за {timeout} секунд") from exc

    raw_response = "".join(_extract_assistant_text(completed.stdout)).strip()
    error_message = _claude_cli_error_message(completed, raw_response)
    if error_message:
        raise ResumeParseError(error_message)
    if not raw_response:
        raise ResumeParseError(completed.stderr.strip() or "Claude CLI не вернул ответ")
    return raw_response


def _run_codex_text(prompt: str, *, timeout: int) -> str:
    output_fd, output_name = tempfile.mkstemp(prefix="career-codex-", suffix=".txt")
    os.close(output_fd)
    output_path = Path(output_name)
    output_path.unlink(missing_ok=True)

    command = [
        CODEX_CLI_PATH,
        "-a",
        "never",
        "-s",
        CODEX_SANDBOX,
        "exec",
        "--skip-git-repo-check",
        "--ignore-user-config",
        "--ignore-rules",
        "--ephemeral",
        "--color",
        "never",
        "-o",
        str(output_path),
    ]
    if CODEX_MODEL:
        command.extend(["-m", CODEX_MODEL])
    command.append("-")

    try:
        completed = subprocess.run(
            command,
            input=prompt,
            capture_output=True,
            text=True,
            timeout=timeout,
            check=False,
        )
    except FileNotFoundError as exc:
        raise ResumeParseError(f"Codex CLI не найден: {CODEX_CLI_PATH}") from exc
    except subprocess.TimeoutExpired as exc:
        content = output_path.read_text(encoding="utf-8").strip() if output_path.exists() else ""
        output_path.unlink(missing_ok=True)
        if content:
            return content
        raise ResumeParseError(f"Codex CLI не ответил за {timeout} секунд") from exc

    content = output_path.read_text(encoding="utf-8").strip() if output_path.exists() else ""
    output_path.unlink(missing_ok=True)
    if completed.returncode != 0:
        error = completed.stderr.strip()[:500] or "без stderr"
        raise ResumeParseError(f"Codex CLI завершился с кодом {completed.returncode}: {error}")
    if not content:
        content = completed.stdout.strip()
    if not content:
        error = completed.stderr.strip()[:500] or "без stderr"
        raise ResumeParseError(f"Codex CLI не вернул ответ: {error}")
    return content


def _slice_lines(lines: list[str], start_line: object, end_line: object) -> str:
    try:
        start = max(1, int(start_line))
        end = min(len(lines), int(end_line))
    except (TypeError, ValueError):
        return ""
    if end < start:
        return ""
    return "\n".join(lines[start - 1 : end])


def _normalize_parsed_resume(payload: dict) -> dict:
    profile = payload.get("profile") if isinstance(payload.get("profile"), dict) else {}
    jobs = payload.get("jobs") if isinstance(payload.get("jobs"), list) else []
    education = payload.get("education") if isinstance(payload.get("education"), list) else []

    return {
        "profile": {
            "full_name": str(profile.get("full_name") or ""),
            "email": str(profile.get("email") or ""),
            "phone": str(profile.get("phone") or ""),
            "location": str(profile.get("location") or ""),
            "linkedin_url": str(profile.get("linkedin_url") or ""),
            "github_url": str(profile.get("github_url") or ""),
            "telegram": str(profile.get("telegram") or ""),
            "summary_raw": str(profile.get("summary_raw") or ""),
        },
        "jobs": [_normalize_job(item) for item in jobs if isinstance(item, dict)],
        "skills_raw": str(payload.get("skills_raw") or ""),
        "education": [_normalize_education(item) for item in education if isinstance(item, dict)],
    }


def _normalize_job(item: dict) -> dict:
    projects = item.get("projects") if isinstance(item.get("projects"), list) else []
    return {
        "company": str(item.get("company") or ""),
        "position": str(item.get("position") or ""),
        "location": str(item.get("location") or ""),
        "start_date": str(item.get("start_date") or ""),
        "end_date": str(item.get("end_date") or ""),
        "is_current": bool(item.get("is_current")),
        "raw_notes": str(item.get("raw_notes") or ""),
        "projects": [_normalize_project(project) for project in projects if isinstance(project, dict)],
    }


def _normalize_project(item: dict) -> dict:
    return {
        "name": str(item.get("name") or ""),
        "raw_description": str(item.get("raw_description") or ""),
        "tech_stack": str(item.get("tech_stack") or ""),
        "results_raw": str(item.get("results_raw") or ""),
        "my_role": str(item.get("my_role") or ""),
        "team_size": str(item.get("team_size") or ""),
    }


def _normalize_education(item: dict) -> dict:
    return {
        "institution": str(item.get("institution") or ""),
        "degree": str(item.get("degree") or ""),
        "field": str(item.get("field") or ""),
        "start_year": str(item.get("start_year") or ""),
        "end_year": str(item.get("end_year") or ""),
        "raw_notes": str(item.get("raw_notes") or ""),
    }


def _build_outline_prompt(raw_text: str) -> str:
    return f"""Ты парсер резюме. Прочитай сырой текст резюме ниже.

Верни ТОЛЬКО валидный JSON без markdown. Это первый быстрый проход: не копируй длинные описания проектов и обязанностей.

Структура JSON:
{{
  "profile": {{
    "full_name": "",
    "email": "",
    "phone": "",
    "location": "",
    "linkedin_url": "",
    "github_url": "",
    "telegram": "",
    "summary_raw": ""
  }},
  "jobs": [
    {{
      "company": "",
      "position": "",
      "location": "",
      "start_date": "",
      "end_date": "",
      "is_current": false,
      "raw_notes": "",
      "start_line": 1,
      "end_line": 1,
      "projects": []
    }}
  ],
  "skills_raw": "",
  "education": [
    {{
      "institution": "",
      "degree": "",
      "field": "",
      "start_year": "",
      "end_year": "",
      "raw_notes": ""
    }}
  ]
}}

Правила:
- jobs[] должен содержать все места работы в порядке как в резюме
- start_line/end_line — диапазон строк этого места работы в файле
- raw_notes в этом проходе: только 1-2 коротких предложения, без длинных списков
- start_date и end_date — строки как в резюме ("март 2021", "2021", "01.2021")
- is_current: true если "по настоящее время", "present", "н.в." и т.д.
- Если поле не найдено — пустая строка, не null

Сырой текст резюме:
{raw_text}
"""


def _build_outline_file_prompt(resume_text_path: Path) -> str:
    return f"""Ты парсер резюме. Прочитай текст резюме из файла:
{resume_text_path}

Верни ТОЛЬКО валидный JSON без markdown. Это первый быстрый проход: не копируй длинные описания проектов и обязанностей.

Структура JSON:
{{
  "profile": {{
    "full_name": "",
    "email": "",
    "phone": "",
    "location": "",
    "linkedin_url": "",
    "github_url": "",
    "telegram": "",
    "summary_raw": ""
  }},
  "jobs": [
    {{
      "company": "",
      "position": "",
      "location": "",
      "start_date": "",
      "end_date": "",
      "is_current": false,
      "raw_notes": "",
      "start_line": 1,
      "end_line": 1,
      "projects": []
    }}
  ],
  "skills_raw": "",
  "education": [
    {{
      "institution": "",
      "degree": "",
      "field": "",
      "start_year": "",
      "end_year": "",
      "raw_notes": ""
    }}
  ]
}}

Правила:
- jobs[] должен содержать все места работы в порядке как в резюме
- start_line/end_line — диапазон строк этого места работы в файле
- raw_notes в этом проходе: только 1-2 коротких предложения, без длинных списков
- start_date и end_date — строки как в резюме ("март 2021", "2021", "01.2021")
- is_current: true если "по настоящее время", "present", "н.в." и т.д.
- Если поле не найдено — пустая строка, не null
"""


def _build_job_prompt(job: dict, job_text: str) -> str:
    company = job.get("company") or ""
    position = job.get("position") or ""
    return f"""Ты парсер одного места работы из резюме.

Компания из первого прохода: {company}
Должность из первого прохода: {position}

Верни ТОЛЬКО валидный JSON без markdown:
{{
  "company": "",
  "position": "",
  "location": "",
  "start_date": "",
  "end_date": "",
  "is_current": false,
  "raw_notes": "",
  "projects": [
    {{
      "name": "",
      "raw_description": "",
      "tech_stack": "",
      "results_raw": "",
      "my_role": "",
      "team_size": ""
    }}
  ]
}}

Правила:
- Сохраняй исходные формулировки из текста, не улучшай стиль.
- Если в тексте есть нумерованные проекты, вынеси их в projects[].
- Достижения конкретного проекта клади в results_raw этого проекта.
- Общие обязанности по месту работы клади в raw_notes.
- Если проектов нет, projects оставь [].
- Если поле не найдено — пустая строка, не null.

Текст места работы:
{job_text}
"""


def _build_codex_resume_prompt(raw_text: str) -> str:
    return f"""Ты парсер резюме. Твоя задача — извлечь из сырого текста резюме структурированные данные.

Верни ТОЛЬКО валидный JSON без markdown, комментариев и текста вокруг JSON.

Структура JSON:
{{
  "profile": {{
    "full_name": "",
    "email": "",
    "phone": "",
    "location": "",
    "linkedin_url": "",
    "github_url": "",
    "telegram": "",
    "summary_raw": ""
  }},
  "jobs": [
    {{
      "company": "",
      "position": "",
      "location": "",
      "start_date": "",
      "end_date": "",
      "is_current": false,
      "raw_notes": "",
      "projects": [
        {{
          "name": "",
          "raw_description": "",
          "tech_stack": "",
          "results_raw": "",
          "my_role": "",
          "team_size": ""
        }}
      ]
    }}
  ],
  "skills_raw": "",
  "education": [
    {{
      "institution": "",
      "degree": "",
      "field": "",
      "start_year": "",
      "end_year": "",
      "raw_notes": ""
    }}
  ]
}}

Правила:
- Сохраняй исходные формулировки из резюме, не улучшай стиль.
- jobs[] должен содержать все места работы в порядке как в резюме.
- Если внутри места работы перечислены проекты, вынеси их в projects[].
- Достижения конкретного проекта клади в results_raw этого проекта.
- Общие обязанности по месту работы клади в raw_notes.
- start_date и end_date — строки как в резюме.
- is_current: true если работа продолжается по настоящее время.
- Если поле не найдено — пустая строка, не null.

Сырой текст резюме:
{raw_text}
"""
